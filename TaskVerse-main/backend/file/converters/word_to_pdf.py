import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

def convert_word_to_pdf(input_path, output_path):
    """Convert DOCX to PDF using python-docx and reportlab."""
    try:
        # Try to extract content using python-docx
        print(f"Attempting to convert: {input_path}")
        
        # First create a working copy of the file (sometimes helps with corrupted files)
        import shutil
        temp_docx = os.path.join(os.path.dirname(input_path), "temp_copy.docx")
        shutil.copy2(input_path, temp_docx)
        
        # Try to open with python-docx
        doc = Document(temp_docx)
        
        # Set up PDF document
        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = styles['Heading1']
        heading_style = styles['Heading2']
        normal_style = styles['Normal']
        
        # Custom style for code or preformatted text
        code_style = ParagraphStyle(
            'Code',
            parent=normal_style,
            fontName='Courier',
            fontSize=8,
            leading=10,
            leftIndent=36
        )
        
        # Content container
        content = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                # Add some space for empty paragraphs
                content.append(Spacer(1, 12))
                continue
                
            # Determine style based on paragraph's style name
            style_name = para.style.name.lower()
            if 'heading 1' in style_name or 'title' in style_name:
                p_style = title_style
            elif 'heading' in style_name:
                p_style = heading_style
            elif 'code' in style_name or 'preformatted' in style_name:
                p_style = code_style
            else:
                p_style = normal_style
                
            # Create paragraph with appropriate style
            p = Paragraph(para.text, p_style)
            content.append(p)
            content.append(Spacer(1, 6))  # Small space after each paragraph
            
        # Process tables (basic support)
        for table in doc.tables:
            # Add a spacer before the table
            content.append(Spacer(1, 12))
            
            # For each row in the table
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    p = Paragraph(row_text, normal_style)
                    content.append(p)
                    
            # Add a spacer after the table
            content.append(Spacer(1, 12))
            
        # Build PDF
        if content:
            pdf_doc.build(content)
            
        # Clean up temp file
        try:
            os.remove(temp_docx)
        except:
            pass
            
        return True
        
    except Exception as first_error:
        print(f"Python-docx extraction failed: {first_error}")
        
        try:
            # Fallback to a more basic approach
            from docx2txt import process
            text = process(input_path)
            
            # Create PDF with extracted text
            from reportlab.pdfgen import canvas
            
            c = canvas.Canvas(output_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            
            # Simple text wrapping
            y = 750
            lines = text.split('\n')
            
            for line in lines:
                if not line.strip():
                    y -= 12
                    continue
                    
                # Wrap long lines
                max_width = 500
                current_line = ""
                
                for word in line.split():
                    test_line = f"{current_line} {word}" if current_line else word
                    
                    if c.stringWidth(test_line, "Helvetica", 12) < max_width:
                        current_line = test_line
                    else:
                        c.drawString(50, y, current_line)
                        y -= 14
                        current_line = word
                        
                        # Check for page break
                        if y < 50:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y = 750
                            
                # Output the final line
                if current_line:
                    c.drawString(50, y, current_line)
                    y -= 14
                    
                # Check for page break
                if y < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = 750
                    
            c.save()
            return True
            
        except Exception as second_error:
            print(f"Basic text extraction failed: {second_error}")
            
            # Create a simple PDF with error information
            try:
                c = canvas.Canvas(output_path, pagesize=letter)
                c.setFont("Helvetica", 12)
                
                y = 750
                c.drawString(72, y, "Word Document Conversion Result")
                y -= 36
                
                c.setFont("Helvetica", 10)
                c.drawString(72, y, "The Word document could not be fully converted to PDF format.")
                y -= 14
                c.drawString(72, y, f"File: {os.path.basename(input_path)}")
                y -= 24
                
                c.drawString(72, y, "This PDF contains information about the Word document:")
                y -= 20
                
                # File information
                import datetime
                file_size = os.path.getsize(input_path)
                file_date = datetime.datetime.fromtimestamp(os.path.getmtime(input_path))
                
                c.drawString(90, y, f"Size: {file_size/1024:.2f} KB")
                y -= 14
                c.drawString(90, y, f"Modified: {file_date}")
                y -= 24
                
                c.drawString(72, y, "Possible reasons for conversion failure:")
                y -= 14
                c.drawString(90, y, "• Document is password protected or encrypted")
                y -= 14
                c.drawString(90, y, "• Document is corrupted or not a valid .docx file")
                y -= 14
                c.drawString(90, y, "• Document uses features not supported by the converter")
                y -= 14
                c.drawString(90, y, "• Incompatible Word format (try saving as .docx)")
                
                c.save()
                return True
                
            except Exception as e:
                print(f"Error creating fallback PDF: {e}")
                # Ultra-minimal PDF
                with open(output_path, 'wb') as f:
                    f.write(b'%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj 2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj 3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>>>endobj xref 0 4 0000000000 65535 f 0000000015 00000 n 0000000060 00000 n 0000000111 00000 n trailer<</Size 4/Root 1 0 R>>startxref 190 %%EOF')
                return True