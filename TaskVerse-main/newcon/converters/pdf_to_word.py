import fitz
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import os

def extract_images_from_page(page, doc, image_counter):
    """Extract and save images from a PDF page"""
    image_list = page.get_images()
    
    for img_index, img in enumerate(image_list):
        try:
            # Get image data
            xref = img[0]
            base_image = page.parent.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Save image to memory
            image_stream = io.BytesIO(image_bytes)
            image = Image.open(image_stream)
            
            # Save image to a temporary file
            temp_img_path = f'temp_img_{image_counter}_{img_index}.png'
            image.save(temp_img_path)
            
            # Add image to document
            doc.add_picture(temp_img_path, width=Inches(6.0))  # Adjust width as needed
            
            # Clean up temporary file
            os.remove(temp_img_path)
            
        except Exception as e:
            print(f"Error processing image: {str(e)}")

def convert_pdf_to_word(input_path, output_path):
    """Convert PDF to Word while preserving formatting and images"""
    pdf_document = fitz.open(input_path)
    doc = Document()
    
    # Set default font and margins
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    image_counter = 0
    
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        
        # Extract images first
        extract_images_from_page(page, doc, image_counter)
        image_counter += 1
        
        # Extract text with formatting information
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    if "spans" in line:
                        text = ""
                        for span in line["spans"]:
                            if span.get("text", "").strip():
                                text += span["text"]
                        
                        if text.strip():
                            paragraph = doc.add_paragraph(text)
                            
                            # Try to preserve alignment
                            if block.get("align", 0) == 1:  # Center aligned
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif block.get("align", 0) == 2:  # Right aligned
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:  # Default to left aligned
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save(output_path)
    pdf_document.close()
