import os
import sys
import logging
import subprocess
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def convert_word_to_pdf(input_path, output_path):
    """Convert Word document to PDF using multiple fallback methods"""
    logger.info(f"Starting Word to PDF conversion: {input_path} -> {output_path}")
    
    # Make sure the output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Method 1: Try using LibreOffice (most compatible cross-platform solution)
    try:
        logger.info("Attempting conversion with LibreOffice")
        
        # LibreOffice command paths to try (add more paths if needed)
        libreoffice_commands = [
            'libreoffice', 'soffice', 
            'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
            'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
            # Add more potential paths here
        ]
        
        command = None
        for cmd in libreoffice_commands:
            try:
                process = subprocess.run([cmd, '--version'], 
                                      stdout=subprocess.PIPE, 
                                      stderr=subprocess.PIPE, 
                                      check=False,
                                      timeout=5)
                if process.returncode == 0:
                    command = cmd
                    logger.info(f"Found LibreOffice at: {cmd}")
                    break
            except Exception:
                continue
        
        if command:
            # Convert using LibreOffice
            logger.info(f"Using LibreOffice command: {command}")
            output_dir = os.path.dirname(output_path)
            
            result = subprocess.run([
                command,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                input_path
            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
            
            logger.info(f"LibreOffice conversion stdout: {result.stdout.decode('utf-8', errors='ignore')}")
            logger.info(f"LibreOffice conversion stderr: {result.stderr.decode('utf-8', errors='ignore')}")
            
            # LibreOffice creates PDF with the same name but .pdf extension
            input_filename = os.path.basename(input_path)
            name_without_ext = os.path.splitext(input_filename)[0]
            temp_output_path = os.path.join(output_dir, f"{name_without_ext}.pdf")
            
            if os.path.exists(temp_output_path):
                # Rename to desired output path if needed
                if temp_output_path != output_path:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    os.rename(temp_output_path, output_path)
                    logger.info(f"Renamed {temp_output_path} to {output_path}")
                
                logger.info("LibreOffice conversion successful")
                return True
            else:
                logger.warning(f"LibreOffice conversion failed, output file not created: {temp_output_path}")
        else:
            logger.info("LibreOffice not found")
    except Exception as e:
        logger.warning(f"LibreOffice conversion failed: {str(e)}")
    
    # Method 2: Try using Microsoft Word via COM (Windows only)
    if sys.platform == "win32":
        try:
            logger.info("Attempting conversion with MS Word COM automation")
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Convert full path to absolute path
            abs_input_path = os.path.abspath(input_path)
            abs_output_path = os.path.abspath(output_path)
            
            logger.info(f"Opening document: {abs_input_path}")
            doc = word.Documents.Open(abs_input_path)
            
            # PDF format constant (17)
            logger.info(f"Saving as PDF: {abs_output_path}")
            doc.SaveAs(abs_output_path, FileFormat=17)
            doc.Close()
            word.Quit()
            
            if os.path.exists(output_path):
                logger.info("MS Word COM conversion successful")
                return True
            else:
                logger.warning("MS Word COM did not create output file")
        except Exception as e:
            logger.warning(f"MS Word COM conversion failed: {str(e)}")
    
    # Method 3: Try using python-docx2pdf fallback
    try:
        logger.info("Attempting conversion with docx2pdf")
        from docx2pdf import convert
        convert(input_path, output_path)
        
        if os.path.exists(output_path):
            logger.info("docx2pdf conversion successful")
            return True
        else:
            logger.warning("docx2pdf didn't create output file")
    except Exception as e:
        logger.warning(f"docx2pdf conversion failed: {str(e)}")
    
    # Method 4: Try using UNOCONV if available
    try:
        logger.info("Attempting conversion with unoconv")
        result = subprocess.run([
            'unoconv', '-f', 'pdf', '-o', output_path, input_path
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        if os.path.exists(output_path):
            logger.info("unoconv conversion successful")
            return True
    except Exception as e:
        logger.warning(f"unoconv conversion failed: {str(e)}")
    
    # Method 5: Final fallback using python-docx and reportlab
    try:
        logger.info("Attempting conversion with python-docx and reportlab")
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        # Open the Word document
        doc = Document(input_path)
        
        # Create a PDF document
        pdf = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        normal_style = styles["Normal"]
        
        # Create story (content)
        story = []
        
        # Process paragraphs
        logger.info(f"Processing {len(doc.paragraphs)} paragraphs")
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                story.append(Spacer(1, 12))
                continue
            
            # Create custom style based on paragraph properties
            style = ParagraphStyle(
                f"ParaStyle{i}",
                parent=normal_style,
                spaceBefore=6,
                spaceAfter=6
            )
            
            # Add alignment
            if para.alignment == 1:  # CENTER
                style.alignment = 1
            elif para.alignment == 2:  # RIGHT
                style.alignment = 2
            elif para.alignment == 3:  # JUSTIFY
                style.alignment = 4
            
            # Create paragraph with text
            p = Paragraph(para.text, style)
            story.append(p)
            
            # Add some space between paragraphs
            story.append(Spacer(1, 6))
        
        # Build the PDF
        logger.info("Building PDF")
        pdf.build(story)
        
        if os.path.exists(output_path):
            logger.info("Simple fallback conversion successful")
            return True
        else:
            logger.error("Failed to create PDF with fallback method")
    
    except Exception as e:
        logger.error(f"All conversion methods failed. Final error: {str(e)}")
        raise Exception(f"Word to PDF conversion failed after trying all methods: {str(e)}")
